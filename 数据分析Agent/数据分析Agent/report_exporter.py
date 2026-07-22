import io
import os
import re
import tempfile
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

_FONT_PATH = r"C:\Windows\Fonts\simhei.ttf"
_FONT_NAME = "SimHei"
_MAX_CHARS_PER_LINE = 58


def _wrap_and_write(pdf, text, font_size=9, line_h=5, indent=0):
    """写入一行文本，超长自动折行，支持左缩进"""
    pdf.set_font(_FONT_NAME, size=font_size)
    available = _MAX_CHARS_PER_LINE - indent
    while len(text) > available:
        chunk = text[:available]
        pdf.cell(0, line_h, " " * indent + chunk, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        text = text[available:]
        indent = 2  # 续行缩进
    if text:
        pdf.cell(0, line_h, " " * indent + text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)


def _render_markdown(pdf, text, font_size=9, line_h=5):
    """将 Markdown 文本渲染为格式化 PDF 内容"""
    for raw_line in text.split("\n"):
        line = raw_line.rstrip()

        # 空行
        if not line.strip():
            pdf.ln(2)
            continue

        # 水平分割线 ---
        if re.match(r"^-{3,}$", line.strip()):
            pdf.ln(1)
            pdf.set_draw_color(180, 180, 180)
            x = pdf.get_x()
            y = pdf.get_y()
            pdf.line(x, y, x + 180, y)
            pdf.ln(3)
            continue

        # 标题 ## / ###
        if line.startswith("### "):
            clean = line[4:].strip()
            pdf.set_font(_FONT_NAME, size=10)
            _wrap_and_write(pdf, clean, font_size=10, line_h=6)
            pdf.ln(1)
            continue
        if line.startswith("## "):
            clean = line[3:].strip()
            pdf.set_font(_FONT_NAME, size=11)
            _wrap_and_write(pdf, clean, font_size=11, line_h=7)
            pdf.ln(1)
            continue
        if line.startswith("# "):
            clean = line[2:].strip()
            pdf.set_font(_FONT_NAME, size=12)
            _wrap_and_write(pdf, clean, font_size=12, line_h=8)
            pdf.ln(2)
            continue

        # 有序列表 1. 2. 3.
        m_ordered = re.match(r"^(\d+)\.\s+(.*)", line.strip())
        if m_ordered:
            num = m_ordered.group(1)
            content = _strip_inline_md(m_ordered.group(2))
            _wrap_and_write(pdf, f"{num}. {content}", font_size=font_size,
                            line_h=line_h, indent=0)
            pdf.ln(1)
            continue

        # 无序列表 - / *
        m_bullet = re.match(r"^[-*]\s+(.*)", line.strip())
        if m_bullet:
            content = _strip_inline_md(m_bullet.group(1))
            _wrap_and_write(pdf, f"- {content}", font_size=font_size,
                            line_h=line_h, indent=2)
            pdf.ln(1)
            continue

        # 普通段落（去除行内 ** __ 等）
        clean = _strip_inline_md(line)
        _wrap_and_write(pdf, clean, font_size=font_size, line_h=line_h)


def _strip_inline_md(text):
    """去除行内 Markdown 符号：**粗体** __斜体__ `代码` 等，保留文本"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


class _PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font(_FONT_NAME, "", _FONT_PATH)

    def header(self):
        self.set_font(_FONT_NAME, size=13)
        self.cell(0, 10, "数据分析报告", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
        self.ln(2)

    def footer(self):
        self.set_y(-15)
        self.set_font(_FONT_NAME, size=8)
        self.cell(0, 10, f"第 {self.page_no()} 页", align="C")

    def section_title(self, text):
        self.set_font(_FONT_NAME, size=11)
        self.cell(0, 8, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.ln(1)


def export_pdf(filename_hint, ai_report=None, stat_summary=None,
               figures=None, chat_history=None):
    """生成 PDF 字节流，各板块可独立开关（传 None 则跳过）"""
    pdf = _PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # 文件标题
    pdf.set_font(_FONT_NAME, size=13)
    pdf.cell(0, 10, f"报告：{filename_hint}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    # 统计摘要（纯文本，直接渲染）
    if stat_summary:
        pdf.section_title("统计摘要")
        _render_markdown(pdf, stat_summary, font_size=9, line_h=5)
        pdf.ln(4)

    # AI 解读（Markdown 渲染）
    if ai_report:
        pdf.section_title("AI 解读")
        _render_markdown(pdf, ai_report, font_size=9, line_h=5)
        pdf.ln(4)

    # 追问对话记录
    if chat_history:
        pdf.section_title("追问对话记录")
        for msg in chat_history:
            role_label = "用户" if msg["role"] == "user" else "AI"
            pdf.set_font(_FONT_NAME, size=9)
            pdf.cell(0, 5, f"【{role_label}】", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            _render_markdown(pdf, msg["content"], font_size=9, line_h=5)
            pdf.ln(3)

    # 图表插入
    if figures:
        pdf.add_page()
        pdf.section_title("数据图表")
        for fig in figures:
            try:
                img_bytes = fig.to_image(format="png", width=700, height=350)
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp.write(img_bytes)
                    tmp_path = tmp.name
                pdf.image(tmp_path, w=180)
                pdf.ln(3)
                os.unlink(tmp_path)
            except Exception as e:
                _render_markdown(pdf, f"[图表生成失败：{e}]", font_size=8)

    return bytes(pdf.output())


# ── Word 导出 ────────────────────────────────────────────────

def _docx_render_markdown(doc, text):
    """将 Markdown 文本写入 Word 文档，保留粗体/标题/列表结构"""
    for raw_line in text.split("\n"):
        line = raw_line.rstrip()

        if not line.strip():
            doc.add_paragraph()
            continue

        if re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("─" * 40)
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        # 有序列表
        m_ordered = re.match(r"^(\d+)\.\s+(.*)", line.strip())
        if m_ordered:
            p = doc.add_paragraph(style="List Number")
            _docx_add_inline(p, m_ordered.group(2))
            continue

        # 无序列表
        m_bullet = re.match(r"^[-*]\s+(.*)", line.strip())
        if m_bullet:
            p = doc.add_paragraph(style="List Bullet")
            _docx_add_inline(p, m_bullet.group(1))
            continue

        # 普通段落
        p = doc.add_paragraph()
        _docx_add_inline(p, line)


def _docx_add_inline(paragraph, text):
    """解析行内 **粗体** 并以 Run 写入段落"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # 去除其他行内符号
            clean = re.sub(r"__(.+?)__", r"\1", part)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            clean = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", clean)
            paragraph.add_run(clean)


def export_word(filename_hint, ai_report=None, stat_summary=None,
                figures=None, chat_history=None):
    """生成 Word (.docx) 字节流，各板块可独立开关"""
    doc = Document()

    # 标题
    title = doc.add_heading("数据分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"报告文件：{filename_hint}")
    doc.add_paragraph()

    # 统计摘要
    if stat_summary:
        doc.add_heading("统计摘要", level=1)
        _docx_render_markdown(doc, stat_summary)
        doc.add_paragraph()

    # AI 解读
    if ai_report:
        doc.add_heading("AI 解读", level=1)
        _docx_render_markdown(doc, ai_report)
        doc.add_paragraph()

    # 追问对话记录
    if chat_history:
        doc.add_heading("追问对话记录", level=1)
        for msg in chat_history:
            role_label = "用户" if msg["role"] == "user" else "AI"
            p = doc.add_paragraph()
            p.add_run(f"【{role_label}】").bold = True
            _docx_render_markdown(doc, msg["content"])
            doc.add_paragraph()

    # 图表插入
    if figures:
        doc.add_heading("数据图表", level=1)
        for fig in figures:
            try:
                img_bytes = fig.to_image(format="png", width=700, height=350)
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp.write(img_bytes)
                    tmp_path = tmp.name
                doc.add_picture(tmp_path, width=Inches(6))
                doc.add_paragraph()
                os.unlink(tmp_path)
            except Exception as e:
                doc.add_paragraph(f"[图表生成失败：{e}]")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
